import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np

# read Excel file
Fin_Statement = pd.ExcelFile("C:\\Users\\bujji143\\Desktop\\GIL\\DILIP WASUDEO.xlsx")
Balance_Sheet = pd.read_excel(Fin_Statement, 'Balance Sheet', header=None)
ProfitAndLossSheet = pd.read_excel(Fin_Statement, 'Income Statement', header=None)

# read Industry excel file
Industry_Statement = pd.ExcelFile("C:\\Users\\bujji143\\Desktop\\GIL\\Average_for_ratios.xlsx")
Balance_Sheet_Industry = pd.read_excel(Industry_Statement, 'Balance Sheet', header=None)
ProfitAndLossSheet_Industry = pd.read_excel(Industry_Statement, 'Income Statement', header=None)
# initializing lists of different ratios
Year_List = list()
Current_Ratio_List = list()
Current_Ratio_Industry_List = list()
Quick_Ratio_List = list()
Quick_Ratio_Industry_List = list()
Return_On_Assets_List = list()
Return_On_Assets_Industry_List = list()
Return_On_Equity_List = list()
Return_On_Equity_Industry_List = list()
Gross_Margin_List = list()
Gross_Margin_Industry_List = list()
Profit_Margin_List = list()
Profit_Margin_Industry_List = list()
Operating_Margin_List = list()
Operating_Margin_Industry_List = list()
Asset_TurnOver_List = list()
Asset_TurnOver_Industry_List = list()
Accounts_Receivable_Turnover_List = list()
Accounts_Receivable_Turnover_Industry_List = list()
Accounts_Payable_Days_List = list()
Accounts_Payable_Days_Industry_List = list()
Average_Days_Sales_List = list()
Average_Days_Sales_Industry_List = list()
Days_Receivable_List = list()
Days_Receivable_Industry_List = list()
Inventory_TurnOver_List = list()
Inventory_TurnOver_Industry_List = list()
Inventory_TurnOver_Period_List = list()
Inventory_TurnOver_Period_Industry_List = list()
Fixed_Asset_TurnOver_List = list()
Fixed_Asset_TurnOver_Industry_List = list()
Working_Capital_TurnOver_List = list()
Working_Capital_TurnOver_Industry_List = list()
Debt_Ratio_List = list()
Debt_ratio_Industry_List = list()
Debt_Equity_Ratio_List = list()
Debt_Equity_Ratio_Industry_List = list()

# Common size analysis
Revenue_List = list()
Revenue_Industry_List = list()
Beginning_Inventory_Percentage_List = list()
Beginning_Inventory_Industry_Percentage_List = list()
Purchases_Percentage_List = list()
Purchases_Industry_Percentage_List = list()
Freight_Expenses_Percentage_List = list()
Freight_Expenses_Industry_Percentage_List = list()
Ending_Inventory_Percentage_List = list()
Ending_Inventory_Industry_Percentage_List = list()
Bill_Of_Materials_Percentage_List = list()
Bill_Of_Materials_Industry_Percentage_List = list()
Labour_Charges_Percentage_List = list()
Labour_Charges_Per_Working_Hour_List = list()
Labour_Charges_Industry_Percentage_List = list()
Labour_Charges_Industry_Per_Working_Hour_List=list()
Sub_Contract_Expenses_Percentage_List = list()
Sub_Contract_Expenses_Industry_Percentage_List = list()
Cost_Of_Goods_Sold_Percentage_List = list()
Cost_Of_Goods_Sold_Percentage_Industry_List = list()
Gross_Profit_Percentage_List = list()
Gross_Profit_Industry_Percentage_List = list()
Business_Development_Expenses_Percentage_List = list()
Business_Development_Expenses_Industry_Percentage_List = list()
Fuel_Expenses_Percentage_List = list()
Fuel_Expenses_Industry_Percentage_List = list()
Conveyance_Expenses_Percentage_List = list()
Conveyance_Expenses_Industry_Percentage_List = list()
Telephone_Expenses_Percentage_List = list()
Telephone_Expenses_Industry_Percentage_List = list()
Selling_And_Admin_Expenses_Percentage_List = list()
Selling_And_Admin_Expenses_Industry_Percentage_List = list()
Electricity_Expenses_Percentage_List = list()
Electricity_Expenses_Industry_Percentage_List = list()
Vehicle_Maintenance_Percentage_List = list()
Vehicle_Maintenance_Industry_Percentage_List = list()
Machine_Maintenance_Percentage_List = list()
Machine_Maintenance_Industry_Percentage_List = list()
Rent_Percentage_List = list()
Rent_Industry_Percentage_List = list()
Consumables_Percentage_List = list()
Consumables_Industry_Percentage_List = list()
Bank_Charges_Percentage_List = list()
Bank_Charges_Industry_Percentage_List = list()
Other_Operating_Expenses_Percentage_List = list()
Other_Operating_Expenses_Industry_Percentage_List = list()
EBITDA_Percentage_List = list()
EBITDA_Industry_Percentage_List = list()
Depreciation_Expense_Percentage_List = list()
Depreciation_Industry_Expense_Percentage_List = list()
Amortization_Expense_Percentage_List = list()
Amortization_Expense_Industry_Percentage_List = list()
Operating_Profit_Percentage_List = list()
Operating_Profit_Industry_Percentage_List = list()
Bank_Interest_Expense_Percentage_List = list()
Bank_Interest_Expense_Industry_Percentage_List = list()
Other_Interest_Expense_Percentage_List = list()
Other_Interest_Expense_Industry_Percentage_List = list()
Total_Interest_Expense_Percentage_List = list()
Total_Interest_Expense_Industry_Percentage_List = list()
Interest_Expense_Per_Working_Hour_List = list()
Interest_Expense_Industry_Per_Working_Hour_List=list()
Other_Expenses_Percentage_List = list()
Other_Expenses_Industry_Percentage_List = list()
Bank_Interest_Income_Percentage_List = list()
Bank_Interest_Income_Industry_Percentage_List = list()
other_income_Percentage_List = list()
other_income_Industry_Percentage_List = list()
Total_Interest_Income_Percentage_List = list()
Total_Interest_Income_Industry_Percentage_List = list()
Net_Income_Before_Taxes_Percentage_List = list()
Net_Income_Before_Taxes_Industry_Percentage_List = list()
Income_Tax_Expense_Percentage_List = list()
Income_Tax_Expense_Industry_Percentage_List = list()
Net_Income_Percentage_List = list()
Net_Income_Industry_Percentage_List = list()

# Asset trend analysis list
Cash_and_Cash_Equivalents_list = list()
Cash_and_Cash_Equivalents_Industry_list = list()
Cash_and_Cash_Equivalents_percentage_change_list = list()
Deposit_list = list()
Deposit_percentage_change_list = list()
Accounts_Receivable_List = list()
Accounts_Receivable_percentage_change_List = list()
Inventory_List = list()
Inventory_Industry_List = list()




Inventory_percentage_change_List = list()
Prepaid_Expenses_List = list()
Prepaid_Expenses_percentage_change_List = list()
Other_Current_Assets_List = list()
Other_Current_Assets_percentage_change_List = list()
Total_Current_Assets_List = list()
Total_Current_Assets_Industry_List = list()
Total_Current_Assets_percentage_change_List = list()
Property_Plant_And_Equipment_percentage_change_List = list()
Other_Long_Term_Assets_percentage_change_List = list()
Intangible_Assets_percentage_change_List = list()
Good_Will_percentage_change_List = list()
Total_Long_Term_Assets_percentage_change_List = list()
Total_Assets_percentage_change_List = list()





number_of_years = 0
for number in range(1, 5):
    Total_Equity = Balance_Sheet[number][39]
    if Total_Equity == 0:
        number_of_years = number_of_years
    else:
        number_of_years = number_of_years + 1
print(number_of_years)
for Col in range(1, number_of_years+1):
    Year = Balance_Sheet[Col][0]
    Year_List.append(Year)
    Current_Assets = Balance_Sheet[Col][8]
    Current_Assets_Industry = Balance_Sheet_Industry[Col][8]
    Current_Liabilities = Balance_Sheet[Col][26]
    Current_Liabilities_Industry = Balance_Sheet_Industry[Col][26]
    # Current_Ratio
    Current_Ratio = Current_Assets/Current_Liabilities
    Current_Ratio_Industry = Current_Assets_Industry/Current_Liabilities_Industry
    # Appending the ratio to the list
    Current_Ratio_List.append(Current_Ratio)
    Current_Ratio_Industry_List.append(Current_Ratio_Industry)
    Inventory = Balance_Sheet[Col][4]
    Quick_Assets = Current_Assets-Inventory
    Inventory_Industry = Balance_Sheet_Industry[Col][4]
    Quick_Assets_Industry = Current_Assets_Industry-Inventory_Industry
    # Quick Ratio
    Quick_Ratio = Quick_Assets/Current_Liabilities
    Quick_Ratio_Industry = Quick_Assets_Industry/Current_Liabilities_Industry
    # Appending the ratio to the list
    Quick_Ratio_List.append(Quick_Ratio)
    Quick_Ratio_Industry_List.append(Quick_Ratio_Industry)
    Net_Income = ProfitAndLossSheet[Col][51]
    Net_Income_Industry = ProfitAndLossSheet_Industry[Col][51]
    Total_Assets = Balance_Sheet[Col][18]
    Total_Assets_Industry = Balance_Sheet_Industry[Col][18]
    # Return On Assets
    Return_On_Assets = (Net_Income/Total_Assets)*100
    Return_On_Assets_Industry = (Net_Income_Industry/Total_Assets_Industry)*100
    # Appending the ratio to the list
    Return_On_Assets_List.append(Return_On_Assets)
    Return_On_Assets_Industry_List.append(Return_On_Assets_Industry)
    Owners_Equity = Balance_Sheet[Col][37]
    Owners_Equity_Industry = Balance_Sheet_Industry[Col][37]
    # Return On Equity
    Return_On_Equity = (Net_Income/Owners_Equity)*100
    Return_On_Equity_Industry = (Net_Income_Industry/Owners_Equity_Industry)*100
    # Appending the ratio to the list
    Return_On_Equity_List.append(Return_On_Equity)
    Return_On_Equity_Industry_List.append(Return_On_Equity_Industry)
    Gross_Profit = ProfitAndLossSheet[Col][11]
    Gross_Profit_Industry = ProfitAndLossSheet_Industry[Col][11]
    Revenue = ProfitAndLossSheet[Col][1]
    Revenue_Industry = ProfitAndLossSheet_Industry[Col][1]
    # Gross Margin
    Gross_Margin = (Gross_Profit/Revenue)*100
    Gross_Margin_Industry = (Gross_Profit_Industry/Revenue_Industry)*100
    # Appending the ratio to the list
    Gross_Margin_List.append(Gross_Margin)
    Gross_Margin_Industry_List.append(Gross_Margin_Industry)
    # Profit Margin
    Profit_Margin = (Net_Income/Revenue)*100
    Profit_Margin_Industry = (Net_Income_Industry/Revenue_Industry)*100
    # Appending the ratio to the list
    Profit_Margin_List.append(Profit_Margin)
    Profit_Margin_Industry_List.append(Profit_Margin_Industry)
    Operating_Income = ProfitAndLossSheet[Col][35]
    Operating_Income_Industry = ProfitAndLossSheet_Industry[Col][35]
    # Operating Margin
    Operating_Margin = (Operating_Income/Revenue)*100
    Operating_Margin_Industry = (Operating_Income_Industry/Revenue_Industry)*100
    # Appending the ratio to the list
    Operating_Margin_List.append(Operating_Margin)
    Operating_Margin_Industry_List.append(Operating_Margin_Industry)
    # Asset Turn Over ratio
    Asset_TurnOver = Revenue/Total_Assets
    Asset_TurnOver_Industry = Revenue_Industry/Total_Assets_Industry
    # Appending the ratio to the list
    Asset_TurnOver_List.append(Asset_TurnOver)
    Asset_TurnOver_Industry_List.append(Asset_TurnOver_Industry)
    Ending_Accounts_Receivable = Balance_Sheet[Col][3]
    Ending_Accounts_Receivable_Industry = Balance_Sheet_Industry[Col][3]
    Beginning_Accounts_Receivable = Balance_Sheet[Col+1][3]
    Beginning_Accounts_Receivable_Industry = Balance_Sheet_Industry[Col+1][3]
    if Beginning_Accounts_Receivable == 0:
        Beginning_Accounts_Receivable = Ending_Accounts_Receivable
    Average_Accounts_Receivable = (Ending_Accounts_Receivable + Beginning_Accounts_Receivable)/2
    if Beginning_Accounts_Receivable_Industry == 0:
        Beginning_Accounts_Receivable_Industry = Ending_Accounts_Receivable_Industry
    Average_Accounts_Receivable_Industry = (Ending_Accounts_Receivable_Industry + Beginning_Accounts_Receivable_Industry)/2
    # Accounts Receivable Turn Over
    Accounts_Receivable_Turnover = Revenue/Average_Accounts_Receivable
    Accounts_receivable_Turnover_Industry = Revenue_Industry/Average_Accounts_Receivable_Industry
    # Appending the ratio to the list
    Accounts_Receivable_Turnover_List.append(Accounts_Receivable_Turnover)
    Accounts_Receivable_Turnover_Industry_List.append(Accounts_receivable_Turnover_Industry)
    # Average Days Sales
    Average_Days_Sales = Revenue/365
    Average_Days_Sales_Industry = Revenue_Industry/365
    # Appending the ratio to the list
    Average_Days_Sales_List.append(Average_Days_Sales)
    Average_Days_Sales_Industry_List.append(Average_Days_Sales_Industry)
    # Days Receivable
    Days_Receivable = Ending_Accounts_Receivable/Average_Days_Sales
    Days_Receivable_Industry = Ending_Accounts_Receivable_Industry/Average_Days_Sales_Industry
    # Appending the ratio to the list
    Days_Receivable_List.append(Days_Receivable)
    Days_Receivable_Industry_List.append(Days_Receivable_Industry)
    Cost_Of_Goods_Sold = ProfitAndLossSheet[Col][10]
    Cost_Of_Goods_Sold_Industry = ProfitAndLossSheet_Industry[Col][10]
    Ending_Inventory = ProfitAndLossSheet[Col][6]
    Ending_Inventory_Industry = Balance_Sheet_Industry[Col][4]
    Beginning_Inventory = ProfitAndLossSheet_Industry[Col][3]
    Beginning_Inventory_Industry = Balance_Sheet_Industry[Col+1][4]
    Average_Inventory = (Ending_Inventory + Beginning_Inventory)/2
    if Beginning_Inventory_Industry == 0:
        Beginning_Inventory_Industry = Ending_Inventory_Industry
    Average_Inventory_Industry = (Ending_Inventory_Industry + Beginning_Inventory_Industry) / 2
    # Inventory Turn Over
    Inventory_TurnOver = Cost_Of_Goods_Sold/Average_Inventory
    Inventory_TurnOver_Industry = Cost_Of_Goods_Sold_Industry/Average_Inventory_Industry
    # Appending the ratio to the list
    Inventory_TurnOver_List.append(Inventory_TurnOver)
    Inventory_TurnOver_Industry_List.append(Inventory_TurnOver_Industry)
    # Inventory Turn Over Period
    Inventory_TurnOver_Period = 365/Inventory_TurnOver
    Inventory_TurnOver_Period_Industry = 365/Inventory_TurnOver_Industry
    # Appending the ratio to the list
    Inventory_TurnOver_Period_List.append(Inventory_TurnOver_Period)
    Inventory_TurnOver_Period_Industry_List.append(Inventory_TurnOver_Period_Industry)
    # Fixed Asset Turn Over
    Fixed_Asset_TurnOver = Cost_Of_Goods_Sold/Total_Assets
    Fixed_Asset_TurnOver_Industry = Cost_Of_Goods_Sold_Industry/Total_Assets_Industry
    # Appending the ratio to the list
    Fixed_Asset_TurnOver_List.append(Fixed_Asset_TurnOver)
    Fixed_Asset_TurnOver_Industry_List.append(Fixed_Asset_TurnOver_Industry)
    # Working Capital Turn Over
    Working_Capital_TurnOver = Revenue/(Current_Assets - Current_Liabilities)
    Working_Capital_TurnOver_Industry = Revenue_Industry/(Current_Assets_Industry-Current_Liabilities_Industry)
    # Appending the ratio to the list
    Working_Capital_TurnOver_List.append(Working_Capital_TurnOver)
    Working_Capital_TurnOver_Industry_List.append(Working_Capital_TurnOver_Industry)
    Total_Liabilities = Balance_Sheet[Col][33]
    Total_Liabilities_Industry = Balance_Sheet_Industry[Col][33]
    # Debt Ratio
    Debt_Ratio = Total_Liabilities/Total_Assets
    Debt_Ratio_Industry = Total_Liabilities_Industry/Total_Assets_Industry
    # Appending the ratio to the list
    Debt_Ratio_List.append(Debt_Ratio)
    Debt_ratio_Industry_List.append(Debt_Ratio_Industry)
    # Debt to Equity Ratio
    Debt_Equity_Ratio = Total_Assets/Owners_Equity
    Debt_Equity_Ratio_Industry = Total_Assets_Industry/Owners_Equity_Industry
    # Appending the ratio to the list
    Debt_Equity_Ratio_List.append(Debt_Equity_Ratio)
    Debt_Equity_Ratio_Industry_List.append(Debt_Equity_Ratio_Industry)

    # Common Size Analysis
    # Appending the revenue to the list
    Revenue_List.append(Revenue)
    Revenue_Industry_List.append(Revenue_Industry)
    # as a % of sales
    Beginning_Inventory_Percentage = (Beginning_Inventory/Revenue)*100
    Beginning_Inventory_Percentage_List.append(Beginning_Inventory_Percentage)

    Beginning_Inventory_Industry_Percentage = (Beginning_Inventory_Industry/Revenue_Industry)*100
    Beginning_Inventory_Industry_Percentage_List.append(Beginning_Inventory_Industry_Percentage)

    Purchases = ProfitAndLossSheet[Col][4]
    Purchases_Percentage = (Purchases/Revenue)*100
    Purchases_Percentage_List.append(Purchases_Percentage)

    Purchases_Industry = ProfitAndLossSheet_Industry[Col][4]
    Purchases_Industry_Percentage = (Purchases_Industry/Revenue_Industry)*100
    Purchases_Industry_Percentage_List.append(Purchases_Industry_Percentage)

    Freight_Expenses = ProfitAndLossSheet[Col][5]
    Freight_Expenses_Percentage = (Freight_Expenses/Revenue)*100
    Freight_Expenses_Percentage_List.append(Freight_Expenses_Percentage)

    Freight_Expenses_Industry = ProfitAndLossSheet_Industry[Col][5]
    Freight_Expenses_Industry_Percentage = (Freight_Expenses_Industry/Revenue_Industry)*100
    Freight_Expenses_Industry_Percentage_List.append(Freight_Expenses_Industry_Percentage)

    Ending_Inventory_Percentage = (Ending_Inventory/Revenue)*100
    Ending_Inventory_Percentage_List.append(Ending_Inventory_Percentage)

    Ending_Inventory_Industry_Percentage = (Ending_Inventory_Industry/Revenue_Industry)*100
    Ending_Inventory_Industry_Percentage_List.append(Ending_Inventory_Industry_Percentage)

    Bill_Of_Materials = ProfitAndLossSheet[Col][7]
    Bill_Of_Materials_Percentage = (Bill_Of_Materials/Revenue)*100
    Bill_Of_Materials_Percentage_List.append(Bill_Of_Materials_Percentage)

    Bill_Of_Materials_Industry = ProfitAndLossSheet_Industry[Col][7]
    Bill_Of_Materials_Industry_Percentage = (Bill_Of_Materials_Industry/Revenue_Industry)*100
    Bill_Of_Materials_Industry_Percentage_List.append(Bill_Of_Materials_Industry_Percentage)

    Labour_Charges = ProfitAndLossSheet[Col][8]
    Labour_Charges_Percentage = (Labour_Charges/Revenue)*100
    Labour_Charges_Percentage_List.append(Labour_Charges_Percentage)

    Labour_Charges_Industry = ProfitAndLossSheet_Industry[Col][8]
    Labour_Charges_Industry_Percentage = (Labour_Charges/Revenue)*100
    Labour_Charges_Industry_Percentage_List.append(Labour_Charges_Industry_Percentage)

    Labour_Charges_Per_Working_Hour = Labour_Charges/2592
    Labour_Charges_Per_Working_Hour_List.append(Labour_Charges_Per_Working_Hour)

    Labour_Charges_Industry_Per_Working_Hour = Labour_Charges_Industry/2592
    Labour_Charges_Industry_Per_Working_Hour_List.append(Labour_Charges_Industry_Per_Working_Hour)

    Sub_Contract_Expenses = ProfitAndLossSheet[Col][9]
    Sub_Contract_Expenses_Percentage = (Sub_Contract_Expenses/Revenue)*100
    Sub_Contract_Expenses_Percentage_List.append(Sub_Contract_Expenses_Percentage)

    Sub_Contract_Expenses_Industry = ProfitAndLossSheet[Col][9]
    Sub_Contract_Expenses_Industry_Percentage = (Sub_Contract_Expenses_Industry / Revenue_Industry) * 100
    Sub_Contract_Expenses_Industry_Percentage_List.append(Sub_Contract_Expenses_Industry_Percentage)

    Cost_Of_Goods_Sold = ProfitAndLossSheet[Col][10]
    Cost_Of_Goods_Sold_Percentage = (Cost_Of_Goods_Sold/Revenue)*100
    Cost_Of_Goods_Sold_Percentage_List.append(Cost_Of_Goods_Sold_Percentage)

    Cost_Of_Goods_Sold_Industry = ProfitAndLossSheet_Industry[Col][10]
    Cost_Of_Goods_Sold_Percentage_Industry = (Cost_Of_Goods_Sold_Industry/Revenue_Industry)*100
    Cost_Of_Goods_Sold_Percentage_Industry_List.append(Cost_Of_Goods_Sold_Percentage_Industry)

    Gross_Profit_Percentage_List.append(Gross_Margin)
    Gross_Profit_Industry_Percentage_List.append(Gross_Margin_Industry)

    Accounts_Payable = Balance_Sheet[Col][20]
    Accounts_Payable_Days = (Accounts_Payable * 365 / Cost_Of_Goods_Sold)
    Accounts_Payable_Days_List.append(Accounts_Payable_Days)

    Accounts_Payable_Industry = Balance_Sheet_Industry[Col][20]
    Accounts_Payable_Days_Industry = (Accounts_Payable_Industry * 365 / Cost_Of_Goods_Sold_Industry)
    Accounts_Payable_Days_Industry_List.append(Accounts_Payable_Days_Industry)

    Business_Development_Expenses = ProfitAndLossSheet[Col][13]
    Business_Development_Expenses_Percentage = (Business_Development_Expenses/Revenue)*100
    Business_Development_Expenses_Percentage_List.append(Business_Development_Expenses_Percentage)

    Business_Development_Expenses_Industry = ProfitAndLossSheet_Industry[Col][13]
    Business_Development_Expenses_Industry_Percentage = (Business_Development_Expenses_Industry / Revenue_Industry) * 100
    Business_Development_Expenses_Industry_Percentage_List.append(Business_Development_Expenses_Industry_Percentage)

    Fuel_Expenses = ProfitAndLossSheet[Col][14]
    Fuel_Expenses_Percentage = (Fuel_Expenses/Revenue)*100
    Fuel_Expenses_Percentage_List.append(Fuel_Expenses_Percentage)


    Fuel_Expenses_Industry = ProfitAndLossSheet_Industry[Col][14]
    Fuel_Expenses_Industry_Percentage = (Fuel_Expenses_Industry / Revenue_Industry) * 100
    Fuel_Expenses_Industry_Percentage_List.append(Fuel_Expenses_Industry_Percentage)


    Conveyance_Expenses = ProfitAndLossSheet[Col][15]
    Conveyance_Expenses_Percentage = (Conveyance_Expenses/Revenue)*100
    Conveyance_Expenses_Percentage_List.append(Conveyance_Expenses_Percentage)

    Conveyance_Expenses_Industry = ProfitAndLossSheet_Industry[Col][15]
    Conveyance_Expenses_Industry_Percentage = (Conveyance_Expenses_Industry / Revenue_Industry) * 100
    Conveyance_Expenses_Industry_Percentage_List.append(Conveyance_Expenses_Industry_Percentage)

    Telephone_Expenses = ProfitAndLossSheet[Col][16]
    Telephone_Expenses_Percentage = (Telephone_Expenses/Revenue)*100
    Telephone_Expenses_Percentage_List.append(Telephone_Expenses_Percentage)

    Telephone_Expenses_Industry = ProfitAndLossSheet_Industry[Col][16]
    Telephone_Expenses_Industry_Percentage = (Telephone_Expenses_Industry / Revenue_Industry) * 100
    Telephone_Expenses_Industry_Percentage_List.append(Telephone_Expenses_Industry_Percentage)

    Selling_And_Admin_Expenses = ProfitAndLossSheet[Col][18]
    Selling_And_Admin_Expenses_Percentage = (Selling_And_Admin_Expenses/Revenue)*100
    Selling_And_Admin_Expenses_Percentage_List.append(Selling_And_Admin_Expenses_Percentage)

    Selling_And_Admin_Expenses_Industry = ProfitAndLossSheet_Industry[Col][18]
    Selling_And_Admin_Expenses_Industry_Percentage = (Selling_And_Admin_Expenses_Industry / Revenue_Industry) * 100
    Selling_And_Admin_Expenses_Industry_Percentage_List.append(Selling_And_Admin_Expenses_Industry_Percentage)

    Electricity_Expenses = ProfitAndLossSheet[Col][20]
    Electricity_Expenses_Percentage = (Electricity_Expenses/Revenue)*100
    Electricity_Expenses_Percentage_List.append(Electricity_Expenses_Percentage)

    Electricity_Expenses_Industry = ProfitAndLossSheet_Industry[Col][20]
    Electricity_Expenses_Industry_Percentage = (Electricity_Expenses_Industry / Revenue_Industry) * 100
    Electricity_Expenses_Industry_Percentage_List.append(Electricity_Expenses_Industry_Percentage)

    Vehicle_Maintenance = ProfitAndLossSheet[Col][21]
    Vehicle_Maintenance_Percentage = (Vehicle_Maintenance/Revenue)*100
    Vehicle_Maintenance_Percentage_List.append(Vehicle_Maintenance_Percentage)

    Vehicle_Maintenance_Industry = ProfitAndLossSheet_Industry[Col][21]
    Vehicle_Maintenance_Industry_Percentage = (Vehicle_Maintenance_Industry / Revenue_Industry) * 100
    Vehicle_Maintenance_Industry_Percentage_List.append(Vehicle_Maintenance_Industry_Percentage)

    Machine_Maintenance = ProfitAndLossSheet[Col][22]
    Machine_Maintenance_Percentage = (Machine_Maintenance/Revenue)*100
    Machine_Maintenance_Percentage_List.append(Machine_Maintenance_Percentage)

    Machine_Maintenance_Industry = ProfitAndLossSheet_Industry[Col][22]
    Machine_Maintenance_Industry_Percentage = (Machine_Maintenance_Industry / Revenue_Industry) * 100
    Machine_Maintenance_Industry_Percentage_List.append(Machine_Maintenance_Industry_Percentage)

    Rent = ProfitAndLossSheet[Col][23]
    Rent_Percentage = (Rent/Revenue)*100
    Rent_Percentage_List.append(Rent_Percentage)

    Rent_Industry = ProfitAndLossSheet_Industry[Col][23]
    Rent_Industry_Percentage = (Rent_Industry / Revenue_Industry) * 100
    Rent_Industry_Percentage_List.append(Rent_Industry_Percentage)

    Consumables = ProfitAndLossSheet[Col][24]
    Consumables_Percentage = (Consumables / Revenue) * 100
    Consumables_Percentage_List.append(Consumables_Percentage)

    Consumables_Industry = ProfitAndLossSheet_Industry[Col][24]
    Consumables_Industry_Percentage = (Consumables_Industry / Revenue_Industry) * 100
    Consumables_Industry_Percentage_List.append(Consumables_Industry_Percentage)

    Bank_Charges = ProfitAndLossSheet[Col][25]
    Bank_Charges_Percentage = (Bank_Charges/Revenue)*100
    Bank_Charges_Percentage_List.append(Bank_Charges_Percentage)

    Bank_Charges_Industry = ProfitAndLossSheet_Industry[Col][25]
    Bank_Charges_Industry_Percentage = (Bank_Charges_Industry / Revenue_Industry) * 100
    Bank_Charges_Industry_Percentage_List.append(Bank_Charges_Industry_Percentage)

    Other_Operating_Expenses = ProfitAndLossSheet[Col][27]
    Other_Operating_Expenses_Percentage = (Other_Operating_Expenses / Revenue) * 100
    Other_Operating_Expenses_Percentage_List.append(Other_Operating_Expenses_Percentage)

    Other_Operating_Expenses_Industry = ProfitAndLossSheet_Industry[Col][27]
    Other_Operating_Expenses_Industry_Percentage = (Other_Operating_Expenses_Industry / Revenue_Industry) * 100
    Other_Operating_Expenses_Industry_Percentage_List.append(Other_Operating_Expenses_Industry_Percentage)

    EBITDA = ProfitAndLossSheet[Col][29]
    EBITDA_Percentage = (EBITDA/Revenue)*100
    EBITDA_Percentage_List.append(EBITDA_Percentage)

    EBITDA_Industry = ProfitAndLossSheet_Industry[Col][29]
    EBITDA_Industry_Percentage = (EBITDA_Industry / Revenue_Industry) * 100
    EBITDA_Industry_Percentage_List.append(EBITDA_Industry_Percentage)

    Depreciation_Expense = ProfitAndLossSheet[Col][31]
    Depreciation_Expense_Percentage = (Depreciation_Expense/Revenue)*100
    Depreciation_Expense_Percentage_List.append(Depreciation_Expense_Percentage)

    Depreciation_Industry_Expense = ProfitAndLossSheet_Industry[Col][31]
    Depreciation_Industry_Expense_Percentage = (Depreciation_Industry_Expense / Revenue_Industry) * 100
    Depreciation_Industry_Expense_Percentage_List.append(Depreciation_Industry_Expense_Percentage)

    Amortization_Expense = ProfitAndLossSheet[Col][32]
    Amortization_Expense_Percentage = (Amortization_Expense / Revenue) * 100
    Amortization_Expense_Percentage_List.append(Amortization_Expense_Percentage)

    Amortization_Expense_Industry = ProfitAndLossSheet_Industry[Col][32]
    Amortization_Expense_Industry_Percentage = (Amortization_Expense_Industry / Revenue_Industry) * 100
    Amortization_Expense_Industry_Percentage_List.append(Amortization_Expense_Industry_Percentage)

    Other_Expenses = ProfitAndLossSheet[Col][33]
    Other_Expenses_Percentage = (Other_Expenses / Revenue)*100
    Other_Expenses_Percentage_List.append(Other_Expenses_Percentage)

    Other_Expenses_Industry = ProfitAndLossSheet_Industry[Col][33]
    Other_Expenses_Industry_Percentage = (Other_Expenses_Industry / Revenue_Industry) * 100
    Other_Expenses_Industry_Percentage_List.append(Other_Expenses_Industry_Percentage)

    Operating_Profit = ProfitAndLossSheet[Col][35]
    Operating_Profit_Percentage = (Operating_Profit/Revenue)*100
    Operating_Profit_Percentage_List.append(Operating_Profit_Percentage)

    Operating_Profit_Industry = ProfitAndLossSheet_Industry[Col][35]
    Operating_Profit_Industry_Percentage = (Operating_Profit_Industry / Revenue_Industry) * 100
    Operating_Profit_Industry_Percentage_List.append(Operating_Profit_Industry_Percentage)

    Bank_Interest_Expense = ProfitAndLossSheet[Col][37]
    Bank_Interest_Expense_Percentage = (Bank_Interest_Expense/Revenue)*100
    Bank_Interest_Expense_Percentage_List.append(Bank_Interest_Expense_Percentage)

    Bank_Interest_Expense_Industry = ProfitAndLossSheet_Industry[Col][37]
    Bank_Interest_Expense_Industry_Percentage = (Bank_Interest_Expense_Industry / Revenue_Industry) * 100
    Bank_Interest_Expense_Industry_Percentage_List.append(Bank_Interest_Expense_Industry_Percentage)

    Other_Interest_Expense = ProfitAndLossSheet[Col][38]
    Other_Interest_Expense_Percentage = (Other_Interest_Expense / Revenue) * 100
    Other_Interest_Expense_Percentage_List.append(Other_Interest_Expense_Percentage)

    Other_Interest_Expense_Industry = ProfitAndLossSheet[Col][38]
    Other_Interest_Expense_Industry_Percentage = (Other_Interest_Expense_Industry / Revenue_Industry) * 100
    Other_Interest_Expense_Industry_Percentage_List.append(Other_Interest_Expense_Industry_Percentage)

    Total_Interest_Expense = ProfitAndLossSheet[Col][40]
    Total_Interest_Expense_Percentage = (Total_Interest_Expense/Revenue)*100
    Total_Interest_Expense_Percentage_List.append(Total_Interest_Expense_Percentage)

    Total_Interest_Expense_Industry = ProfitAndLossSheet_Industry[Col][40]
    Total_Interest_Expense_Industry_Percentage = (Total_Interest_Expense_Industry / Revenue_Industry) * 100
    Total_Interest_Expense_Industry_Percentage_List.append(Total_Interest_Expense_Industry_Percentage)

    Interest_Expense_Per_Working_Hour = Total_Interest_Expense/2592
    Interest_Expense_Per_Working_Hour_List.append(Interest_Expense_Per_Working_Hour)

    Interest_Expense_Industry_Per_Working_Hour = Total_Interest_Expense_Industry / 2592
    Interest_Expense_Industry_Per_Working_Hour_List.append(Interest_Expense_Industry_Per_Working_Hour)

    Bank_Interest_Income = ProfitAndLossSheet[Col][42]
    Bank_Interest_Income_Percentage = (Bank_Interest_Income/Revenue)*100
    Bank_Interest_Income_Percentage_List.append(Bank_Interest_Income_Percentage)

    Bank_Interest_Income_Industry = ProfitAndLossSheet_Industry[Col][42]
    Bank_Interest_Income_Industry_Percentage = (Bank_Interest_Income_Industry / Revenue_Industry) * 100
    Bank_Interest_Income_Industry_Percentage_List.append(Bank_Interest_Income_Industry_Percentage)

    other_income = ProfitAndLossSheet[Col][43]
    other_income_Percentage = (other_income/Revenue)*100
    other_income_Percentage_List.append(other_income_Percentage)

    other_income_Industry = ProfitAndLossSheet_Industry[Col][43]
    other_income_Industry_Percentage = (other_income_Industry / Revenue_Industry) * 100
    other_income_Industry_Percentage_List.append(other_income_Industry_Percentage)

    Total_Interest_Income = ProfitAndLossSheet[Col][45]
    Total_Interest_Income_Percentage = (Total_Interest_Income / Revenue) * 100
    Total_Interest_Income_Percentage_List.append(Total_Interest_Income_Percentage)

    Total_Interest_Income_Industry = ProfitAndLossSheet_Industry[Col][45]
    Total_Interest_Income_Industry_Percentage = (Total_Interest_Income_Industry / Revenue) * 100
    Total_Interest_Income_Industry_Percentage_List.append(Total_Interest_Income_Industry_Percentage)

    Net_Income_Before_Taxes = ProfitAndLossSheet[Col][47]
    Net_Income_Before_Taxes_Percentage = (Net_Income_Before_Taxes/Revenue)*100
    Net_Income_Before_Taxes_Percentage_List.append(Net_Income_Before_Taxes_Percentage)

    Net_Income_Before_Taxes_Industry = ProfitAndLossSheet_Industry[Col][47]
    Net_Income_Before_Taxes_Industry_Percentage = (Net_Income_Before_Taxes_Industry / Revenue_Industry) * 100
    Net_Income_Before_Taxes_Industry_Percentage_List.append(Net_Income_Before_Taxes_Industry_Percentage)

    Income_Tax_Expense = ProfitAndLossSheet[Col][49]
    Income_Tax_Expense_Percentage = (Income_Tax_Expense/Revenue)*100
    Income_Tax_Expense_Percentage_List.append(Income_Tax_Expense_Percentage)

    Income_Tax_Expense_Industry = ProfitAndLossSheet_Industry[Col][49]
    Income_Tax_Expense_Industry_Percentage = (Income_Tax_Expense_Industry / Revenue_Industry) * 100
    Income_Tax_Expense_Industry_Percentage_List.append(Income_Tax_Expense_Industry_Percentage)

    Net_Income = ProfitAndLossSheet[Col][51]
    Net_Income_Percentage = (Net_Income/Revenue)*100
    Net_Income_Percentage_List.append(Net_Income_Percentage)

    Net_Income_Industry = ProfitAndLossSheet_Industry[Col][51]
    Net_Income_Industry_Percentage = (Net_Income_Industry / Revenue_Industry) * 100
    Net_Income_Industry_Percentage_List.append(Net_Income_Industry_Percentage)

document = Document()

# Assets Trend Analysis
for i in range(1, number_of_years+1):
    Cash_and_Cash_Equivalents = Balance_Sheet[Col][1]
    Cash_and_Cash_Equivalents_list.append(Cash_and_Cash_Equivalents)

    Deposit = Balance_Sheet[Col][2]
    Deposit_list.append(Deposit)

    Accounts_Receivable = Balance_Sheet[Col][3]
    Accounts_Receivable_List.append(Accounts_Receivable)

    Inventory = Balance_Sheet[Col][4]
    Inventory_List.append(Inventory)

    Prepaid_Expenses = Balance_Sheet[Col][5]
    Prepaid_Expenses_List.append(Prepaid_Expenses)

    Other_Current_Assets = Balance_Sheet[Col][6]
    Other_Current_Assets_List.append(Other_Current_Assets)

    Total_Current_Assets = Balance_Sheet[Col][8]
    Total_Current_Assets_List.append(Other_Current_Assets)

document.add_heading("Liquidity Ratios")
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ratio Name'
for i in range(len(Year_List)):
    hdr_cells[i + 1].text = str(Year_List[i])
#hdr_cells[1].text = str(Year_List[0])
#hdr_cells[2].text = str(Year_List[1])
#hdr_cells[3].text = str(Year_List[2])
row_cells = table.add_row().cells
row_cells[0].text = 'Current Ratio'
for i in range(len(Year_List)):
    row_cells[i+1].text = str(round(Current_Ratio_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Quick Ratio'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Quick_Ratio_List[i], 2))
row_cells = table.add_row().cells

plt.plot(Year_List, Current_Ratio_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Current_Ratio_Industry_List,linestyle='dashed',label="Industry standard")
#plt.plot(Year_List, Current_Ratio_List)
plt.xlabel('Year')
plt.title('Current ratio')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Current ratio')
document.add_picture('Current ratio.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Quick_Ratio_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Quick_Ratio_Industry_List,linestyle='dashed',label='Industry standard')
#plt.plot(Year_List, Quick_Ratio_List)
plt.xlabel('Year')
plt.title('Quick Ratio')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Quick ratio')
document.add_picture('Quick ratio.png', width=Inches(4))
plt.close()

document.add_heading("Profitability Ratios")
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ratio Name'
for i in range(len(Year_List)):
    hdr_cells[i + 1].text = str(Year_List[i])
#hdr_cells[1].text = str(Year_List[0])
#hdr_cells[2].text = str(Year_List[1])
#hdr_cells[3].text = str(Year_List[2])
row_cells = table.add_row().cells
row_cells[0].text = 'Return On Assets'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Return_On_Assets_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Return On Equity'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Return_On_Equity_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Gross Margin'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Gross_Margin_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Profit Margin'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Profit_Margin_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Operating Margin'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Operating_Margin_List[i], 2))

plt.plot(Year_List, Return_On_Assets_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Return_On_Assets_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Return on Assets ')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Return on Assets')
document.add_picture('Return on Assets.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Return_On_Equity_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Return_On_Equity_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Return On Equity ')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Return on Equity')
document.add_picture('Return on Equity.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Gross_Margin_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Gross_Margin_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Gross Margin ')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Gross Margin')
document.add_picture('Gross Margin.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Profit_Margin_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Profit_Margin_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Profit Margin')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Profit Margin')
document.add_picture('Profit Margin.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Operating_Margin_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Operating_Margin_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Operating Margin')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Operating Margin')
document.add_picture('Operating Margin.png', width=Inches(4))
plt.close()

document.add_heading("Activity Turn over Ratios")
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ratio Name'
for i in range(len(Year_List)):
    hdr_cells[i + 1].text = str(Year_List[i])
#hdr_cells[1].text = str(Year_List[0])
#hdr_cells[2].text = str(Year_List[1])
#hdr_cells[3].text = str(Year_List[2])
row_cells = table.add_row().cells
row_cells[0].text = 'Asset Turn Over'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Asset_TurnOver_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Accounts Receivable Turnover'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Accounts_Receivable_Turnover_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Average Days Sales'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Average_Days_Sales_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Days Receivable'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Days_Receivable_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Days Payable'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Accounts_Payable_Days_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Inventory Turnover'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Inventory_TurnOver_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Inventory Turnover Period'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Inventory_TurnOver_Period_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Working Capital Turnover'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Working_Capital_TurnOver_List[i], 2))

plt.plot(Year_List, Asset_TurnOver_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Asset_TurnOver_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Asset TurnOver')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Asset TurnOver')
document.add_picture('Asset TurnOver.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Average_Days_Sales_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Average_Days_Sales_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Average Days Sales')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Average Days Sales')
document.add_picture('Average Days Sales.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Days_Receivable_List, linestyle='solid',label='Company standard')
#plt.plot(Year_List, Days_Receivable_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Days Receivables')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Days Receivables')
document.add_picture('Days Receivables.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Accounts_Payable_Days_List, linestyle='solid',label='Company standard')
#plt.plot(Year_List, Accounts_Payable_Days_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Days Payable')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Days Payable')
document.add_picture('Days Payable.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Inventory_TurnOver_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Inventory_TurnOver_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Inventory')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Inventory Turn Over')
document.add_picture('Inventory Turn Over.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Inventory_TurnOver_Period_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Inventory_TurnOver_Period_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Inventory Turn Over Period')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Inventory Turn Over Period')
document.add_picture('Inventory Turn Over Period.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Working_Capital_TurnOver_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Working_Capital_TurnOver_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Working Capital Turnover')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Inventory Turn Over Period')
document.add_picture('Inventory Turn Over Period.png', width=Inches(4))
plt.close()

document.add_heading("Solvency Ratios")
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ratio Name'
for i in range(len(Year_List)):
    hdr_cells[i + 1].text = str(Year_List[i])
#hdr_cells[1].text = str(Year_List[0])
#hdr_cells[2].text = str(Year_List[1])
#hdr_cells[3].text = str(Year_List[2])
row_cells = table.add_row().cells
row_cells[0].text = 'Debt Ratio'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Debt_Ratio_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Debt To Equity Ratio'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Debt_Equity_Ratio_List[i], 2))

plt.plot(Year_List, Debt_Ratio_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Debt_ratio_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Debt Ratio')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Debt Ratio')
document.add_picture('Debt Ratio.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Debt_Equity_Ratio_List, linestyle='solid',label='Company standard')
plt.plot(Year_List, Debt_Equity_Ratio_Industry_List,linestyle='dashed',label='Industry standard')
plt.xlabel('Year')
plt.title('Debt Equity Ratio')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Debt Equity Ratio')
document.add_picture('Debt Equity Ratio.png', width=Inches(4))
plt.close()

document.add_heading(" Common Size Analysis of Income Statements ( % of Sales)")
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ratio Name'
for i in range(len(Year_List)):
    hdr_cells[i + 1].text = str(Year_List[i])
#hdr_cells[1].text = str(Year_List[0])
#hdr_cells[2].text = str(Year_List[1])
#hdr_cells[3].text = str(Year_List[2])
row_cells = table.add_row().cells
row_cells[0].text = 'Revenue'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Revenue_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Beginning Inventory '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Beginning_Inventory_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Purchases '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Purchases_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Freight Expenses'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Freight_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Ending Inventory '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Ending_Inventory_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Bill Of Materials '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Bill_Of_Materials_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Labour Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Labour_Charges_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Labour Charges Per Working Hour'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Labour_Charges_Per_Working_Hour_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Sub Contract Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Sub_Contract_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Cost Of Goods Sold '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Cost_Of_Goods_Sold_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Business Development Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Business_Development_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Fuel Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Fuel_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Conveyance Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Conveyance_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Telephone Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Telephone_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Selling & Admin Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Selling_And_Admin_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Electricity Charges '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Electricity_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Vehicle Maintenance '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Vehicle_Maintenance_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Machine Maintenance '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Machine_Maintenance_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Rent '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Rent_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Consumables '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Consumables_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Bank Charges '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Bank_Charges_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Other Operating Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Other_Operating_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'EBITDA '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(EBITDA_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Depreciation '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Depreciation_Expense_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Amortization '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Amortization_Expense_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Other Expenses '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Other_Expenses_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Bank Interest Expense '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Bank_Interest_Expense_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Other Interest Expense '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Bank_Interest_Expense_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Total Interest Expense '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Total_Interest_Expense_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Interest Expense Per Working Hour'
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Interest_Expense_Per_Working_Hour_List[i], 2))

row_cells = table.add_row().cells
row_cells[0].text = 'Bank Interest Income '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Bank_Interest_Income_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Other Income '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(other_income_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Income(Bank Interest Income+Other  Income) '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Total_Interest_Income_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Net Income Before Taxes  '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Net_Income_Before_Taxes_Percentage_List[i], 2))
row_cells = table.add_row().cells
row_cells[0].text = 'Income Tax Expense '
for i in range(len(Year_List)):
    row_cells[i + 1].text = str(round(Income_Tax_Expense_Percentage_List[i], 2))

plt.plot(Year_List, Revenue_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Revenue_Industry_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Revenue ( 1e9=10^9=1000000000 )')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Revenue')
document.add_picture('Revenue.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Beginning_Inventory_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Beginning_Inventory_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Beginning Inventory')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Beginning Inventory')
document.add_picture('Beginning Inventory.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Purchases_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Purchases_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Purchases')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Purchases')
document.add_picture('Purchases.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Freight_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Freight_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Freight Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Freight Expenses')
document.add_picture('Freight Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Ending_Inventory_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Ending_Inventory_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Ending Inventory')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Ending Inventory')
document.add_picture('Ending Inventory.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Bill_Of_Materials_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Bill_Of_Materials_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Bill Of Materials')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Bill Of Materials')
document.add_picture('Bill Of materials.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Labour_Charges_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Labour_Charges_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Labour Charges')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Labour charges')
document.add_picture('Labour Charges.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Labour_Charges_Per_Working_Hour_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Labour_Charges_Industry_Per_Working_Hour_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Labour Charges Per Working Hour')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Labour charges Per Working Hour')
document.add_picture('Labour Charges Per Working Hour.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Sub_Contract_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Sub_Contract_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Sub Contract Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Sub Contract Expenses')
document.add_picture('Sub Contract Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Cost_Of_Goods_Sold_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Cost_Of_Goods_Sold_Percentage_Industry_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Cost Of Goods Sold')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Cost of Goods Sold')
document.add_picture('Cost of Goods Sold.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Gross_Profit_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Gross_Profit_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Gross Profit as a % of sales')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Gross Profit as a %')
document.add_picture('Gross Profit as a %.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Business_Development_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Business_Development_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Business Development Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Business Development Expenses')
document.add_picture('Business Development Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Fuel_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Fuel_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Fuel Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Fuel Expenses')
document.add_picture('Fuel Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Conveyance_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Conveyance_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Conveyance Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Conveyance Expenses')
document.add_picture('Conveyance Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Telephone_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Telephone_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Telephone Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Telephone Expenses')
document.add_picture('Telephone Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Selling_And_Admin_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Selling_And_Admin_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Selling & Admin Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Selling & Admin Expenses')
document.add_picture('Selling & Admin Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Electricity_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Electricity_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Electricity Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Electricity Expenses')
document.add_picture('Electricity Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Vehicle_Maintenance_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Vehicle_Maintenance_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Vehicle Maintenance')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Vehicle Maintenance')
document.add_picture('Vehicle Maintenance.png', width=Inches(4))
plt.close()

plt.plot(Year_List,Machine_Maintenance_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Machine_Maintenance_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Machine Maintenance')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Machine Maintenance')
document.add_picture('Machine Maintenance.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Rent_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Rent_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Rent')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Rent')
document.add_picture('Rent.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Consumables_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Consumables_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Consumables')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Consumables as a %')
document.add_picture('Consumables as a %.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Bank_Charges_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Bank_Charges_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Bank Charges')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Bank Charges as a %')
document.add_picture('Bank Charges as a %.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Other_Operating_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Other_Operating_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Other Operating Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Other Operating Expenses')
document.add_picture('Other Operating Expenses.png', width=Inches(4))
plt.close()

plt.plot(Year_List, EBITDA_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, EBITDA_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('EBITDA as a % of sales')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('EBITDA as a % of sales')
document.add_picture('EBITDA as a % of sales.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Depreciation_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Depreciation_Industry_Expense_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Depreciation Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Depreciation Expense')
document.add_picture('Depreciation Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Amortization_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Amortization_Expense_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Amortization Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Amortization Expense')
document.add_picture('Amortization Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Other_Expenses_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Other_Expenses_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Other Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Other Expense')
document.add_picture('Other Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Operating_Profit_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Operating_Profit_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Operating profit')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Operating Profit')
document.add_picture('Operating Profit.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Bank_Interest_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Bank_Interest_Expense_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Bank Interest Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Bank Interest Expense')
document.add_picture('Bank Interest Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Other_Interest_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Other_Interest_Expense_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Other Interest Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Other Interest Expense')
document.add_picture('Other Interest Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Total_Interest_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Total_Interest_Expense_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Total Interest Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Total Interest Expense')
document.add_picture('Total Interest Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Interest_Expense_Per_Working_Hour_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Interest_Expense_Industry_Per_Working_Hour_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Total Interest Expense')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Total Interest Expense')
document.add_picture('Total Interest Expense.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Bank_Interest_Income_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Bank_Interest_Income_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Bank Interest Income')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Bank Interest Income')
document.add_picture('Bank Interest Income.png', width=Inches(4))
plt.close()

plt.plot(Year_List, other_income_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, other_income_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Other Income')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Other Income')
document.add_picture('Other Income.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Total_Interest_Income_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Total_Interest_Income_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Income(Bank Interest Income+Other  Income)')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Income(Bank Interest Income+Other  Income)')
document.add_picture('Income(Bank Interest Income+Other  Income).png', width=Inches(4))
plt.close()

plt.plot(Year_List, Net_Income_Before_Taxes_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Net_Income_Before_Taxes_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Net Income Before Taxes')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Net Income Before Taxes as a %')
document.add_picture('Net Income Before Taxes as a %.png', width=Inches(4))
plt.close()

plt.plot(Year_List, Income_Tax_Expense_Percentage_List,linestyle='solid',label="Company standard")
plt.plot(Year_List, Income_Tax_Expense_Industry_Percentage_List,linestyle='dashed',label="Industry standard")
plt.xlabel('Year')
plt.title('Income Tax Expenses')
plt.xticks(np.arange(min(Year_List), max(Year_List)+1, 1.0))
plt.legend(loc='best')
plt.savefig('Income Tax Expenses as a %')
document.add_picture('Income Tax Expenses as a %.png', width=Inches(4))
plt.close()
document.save('Ratios.docx')

print("OverOver")

